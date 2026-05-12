"""DOCX parser using python-docx — paragraph + table support."""
from __future__ import annotations

import io
from typing import Iterable

from src.ingestion.interfaces import Block, Parser, RawDocument
from src.ingestion.registry import register_parser


class DocxParser:
    name = "docx"
    mime_types = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    extensions = ("docx",)

    def parse(self, doc: RawDocument) -> Iterable[Block]:
        from docx import Document  # python-docx

        f = io.BytesIO(doc.content)
        document = Document(f)

        current_section: str | None = None
        buffer: list[str] = []

        def flush() -> Iterable[Block]:
            if buffer:
                txt = "\n".join(buffer).strip()
                if txt:
                    yield Block(text=txt, section=current_section)
                buffer.clear()

        for para in document.paragraphs:
            text = (para.text or "").strip()
            if not text:
                continue
            style = (para.style.name or "").lower() if para.style else ""
            if style.startswith("heading"):
                yield from flush()
                current_section = text
                continue
            buffer.append(text)
        yield from flush()

        # Tables — emit one block per table
        for ti, table in enumerate(document.tables, start=1):
            rows: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            body = "\n".join(rows).strip()
            if body:
                yield Block(text=body, section=f"Table {ti}", extra={"is_table": True})


register_parser(DocxParser())
